# create_dataset1_docx.py

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_dataset1_docx():
    # Initialize the Document
    doc = Document()

    # Add Title
    doc.add_heading('Dataset 1 Analysis', 0)

    # Introduction
    doc.add_heading('Introduction', level=1)
    intro_paragraph = (
        "This document presents an analysis of the Employee dataset, aiming to predict employee turnover using various "
        "machine learning models. The analysis encompasses data preprocessing, model training, evaluation, and visualization "
        "of learning curves to assess model performance."
    )
    doc.add_paragraph(intro_paragraph)

    # Data Description
    doc.add_heading('Data Description', level=1)
    data_desc_paragraph = (
        "The **Employee** dataset contains information about employees, including demographic details, job-related features, "
        "and whether an employee left the company (`LeaveOrNot`). Below is an overview of the dataset's structure:"
    )
    doc.add_paragraph(data_desc_paragraph)

    # Columns List
    columns = [
        "EmployeeID: Unique identifier for each employee.",
        "Age: Age of the employee.",
        "Department: Department where the employee works.",
        "Education: Education level of the employee.",
        "JobRole: Role of the employee within the department.",
        "MonthlyIncome: Monthly income of the employee.",
        "TotalWorkingYears: Total years of working experience.",
        "EnvironmentSatisfaction: Satisfaction level with the work environment.",
        "JobSatisfaction: Satisfaction level with the job role.",
        "WorkLifeBalance: Work-life balance satisfaction level.",
        "LeaveOrNot: Target variable indicating if the employee left the company (1) or not (0)."
    ]

    for col in columns:
        doc.add_paragraph(f"- **{col}**")

    # Methodology
    doc.add_heading('Methodology', level=1)
    methodology_intro = (
        "The analysis followed a structured approach comprising data preprocessing, model training, evaluation, "
        "and visualization."
    )
    doc.add_paragraph(methodology_intro)

    # Data Preprocessing Section
    doc.add_heading('1. Data Preprocessing', level=2)
    preprocessing_steps = [
        "**Loading Data:** The dataset was loaded using pandas.",
        "**Handling Categorical Variables:** Categorical columns (`Department`, `Education`, `JobRole`) were label encoded to convert them into numerical format.",
        "**Feature Selection:** The target variable `LeaveOrNot` was separated from the feature set.",
        "**Data Splitting:** The data was split into training and testing sets with a test size of 10%.",
        "**Feature Scaling:** Features were standardized using `StandardScaler` to ensure that all features contribute equally to the model performance."
    ]

    for step in preprocessing_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Model Training Section
    doc.add_heading('2. Model Training', level=2)
    training_steps = (
        "Four different classification models were trained to predict employee turnover:\n\n"
        "- **Decision Tree Classifier**\n"
        "- **K-Nearest Neighbors (KNN) Classifier**\n"
        "- **Logistic Regression**\n"
        "- **Perceptron**\n\n"
        "Each model was trained on the standardized training data and evaluated on the test set."
    )
    doc.add_paragraph(training_steps)

    # Evaluation Metrics Section
    doc.add_heading('3. Evaluation Metrics', level=2)
    eval_steps = (
        "- **Accuracy Score:** Measures the proportion of correctly predicted instances.\n"
        "- **Classification Report:** Provides precision, recall, F1-score, and support for each class.\n"
        "- **Learning Curves:** Visual representations of training and cross-validation scores against varying training sizes to assess model performance and potential overfitting."
    )
    doc.add_paragraph(eval_steps, style='List Bullet')

    # Results Section
    doc.add_heading('Results', level=1)

    # Model Performance Table
    doc.add_heading('Model Performance', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Accuracy'

    # Sample Data - Replace these values with actual results
    model_results = {
        'Decision Tree': '0.8500',
        'K-Nearest Neighbors': '0.8200',
        'Logistic Regression': '0.8300',
        'Perceptron': '0.8000'
    }

    for model, accuracy in model_results.items():
        row_cells = table.add_row().cells
        row_cells[0].text = model
        row_cells[1].text = accuracy

    # Classification Reports
    doc.add_paragraph('**Classification Reports:**', style='Intense Quote')

    classification_reports = {
        'Decision Tree': """
                  precision    recall  f1-score   support

               0       0.85      0.87      0.86       100
               1       0.84      0.81      0.83        50

        accuracy                           0.85       150
       macro avg       0.85      0.84      0.85       150
    weighted avg       0.85      0.85      0.85       150
    """,
        'K-Nearest Neighbors': """
                  precision    recall  f1-score   support

               0       0.82      0.85      0.83       100
               1       0.80      0.78      0.79        50

        accuracy                           0.82       150
       macro avg       0.81      0.82      0.81       150
    weighted avg       0.82      0.82      0.82       150
    """,
        'Logistic Regression': """
                  precision    recall  f1-score   support

               0       0.83      0.85      0.84       100
               1       0.82      0.80      0.81        50

        accuracy                           0.83       150
       macro avg       0.83      0.83      0.83       150
    weighted avg       0.83      0.83      0.83       150
    """,
        'Perceptron': """
                  precision    recall  f1-score   support

               0       0.80      0.82      0.81       100
               1       0.78      0.76      0.77        50

        accuracy                           0.80       150
       macro avg       0.79      0.79      0.79       150
    weighted avg       0.80      0.80      0.80       150
    """
    }

    for model, report in classification_reports.items():
        doc.add_heading(model, level=3)
        # Use Courier New font to emulate preformatted text
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(report)
        run.font.name = 'Courier New'
        # Ensure the font name is correctly set in the XML
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        paragraph.paragraph_format.space_before = Inches(0.1)

    # Learning Curves Section
    doc.add_heading('Learning Curves', level=2)
    learning_curves_intro = (
        "The learning curves for each model illustrate the training and cross-validation scores over varying training sizes. "
        "These plots help in understanding the model's ability to generalize."
    )
    doc.add_paragraph(learning_curves_intro)

    # List of Models and Image Paths
    models_images = {
        'Decision Tree Learning Curve': 'images/tuned dataset1.png',
        'K-Nearest Neighbors Learning Curve': 'images/tuned dataset1.png',
        'Logistic Regression Learning Curve': 'images/tuned dataset1.png',
        'Perceptron Learning Curve': 'images/tuned dataset1.png'
    }

    for idx, (title, img_path) in enumerate(models_images.items(), start=1):
        doc.add_paragraph(f"- **{title}**")
        
        # Placeholder for Image
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Insert placeholder if image does not exist
            placeholder = doc.add_paragraph()
            run = placeholder.add_run("[Image Placeholder]")
            run.bold = True
            placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Figure Caption
        doc.add_paragraph(f"*Figure {idx}: {title}.*", style='Intense Quote')
        doc.add_paragraph()  # Add an empty paragraph for spacing

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion_paragraph = (
        "The analysis demonstrates the performance of various machine learning models in predicting employee turnover. "
        "The Decision Tree Classifier achieved the highest accuracy, closely followed by Logistic Regression and K-Nearest Neighbors. "
        "The Perceptron model showed the least accuracy among the tested models.\n\n"
        "Learning curves indicate that the models generally perform well with the given data, with no significant signs of overfitting. "
        "Future work could involve hyperparameter tuning, feature engineering, and exploring more advanced models to further enhance prediction accuracy."
    )
    doc.add_paragraph(conclusion_paragraph)

    # Notes Section
    doc.add_heading('Notes', level=1)
    notes = [
        "**Figures:** Ensure that the PNG images (e.g., learning curves) are saved in the specified paths (`images/`) and inserted into the document accordingly.",
        "**Tables and Reports:** The classification reports and accuracy scores should reflect the actual results obtained from running your Python script.",
        "**Customization:** Feel free to adjust the content, add more sections (e.g., Feature Importance), or include additional visualizations based on your specific analysis and findings."
    ]

    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

    # Save the Document
    doc.save('dataset 1.docx')
    print("dataset1.docx has been created successfully in the 'CSC272-ML-Project1' directory.")

if __name__ == "__main__":
    create_dataset1_docx()